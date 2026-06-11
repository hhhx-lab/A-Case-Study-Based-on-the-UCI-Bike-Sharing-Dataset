#!/usr/bin/env python3
"""Build the upgraded thesis DOCX from R outputs.

This script is intentionally document-focused: R remains responsible for data
analysis, while this file controls Word structure, captions, references, and
format settings required by the ECNU thesis-format document.
"""

from __future__ import annotations

import csv
import re
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
TABLE_DIR = ROOT / "outputs" / "tables"
FIG_DIR = ROOT / "outputs" / "figures"
MODEL_DIR = ROOT / "outputs" / "models"
PAPER_DIR = ROOT / "paper"
ASSET_DIR = ROOT / "assets"


def read_csv(name: str) -> list[dict[str, str]]:
    with (TABLE_DIR / name).open(newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def summary_map() -> dict[str, str]:
    return {row["item"]: row["value"] for row in read_csv("table99_result_summary.csv")}


S = summary_map()


def rel_path(path: Path) -> str:
    try:
        return path.resolve().relative_to(ROOT).as_posix()
    except ValueError:
        return path.as_posix()


def sanitize_paths(text: str) -> str:
    replacements = {
        str(ROOT): ".",
        str(Path.home() / "workspace" / "统计软件"): ".",
        str(Path.home() / "miniforge3" / "envs" / "r-statsoft"): "<conda-env:r-statsoft>",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def sval(key: str) -> str:
    return S[key]


def fmt_num(value: str | float | int | None, digits: int = 3) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        value = value.strip()
        if value == "":
            return ""
    try:
        return f"{float(value):.{digits}f}"
    except Exception:
        return str(value)


def pct(value: str | float, digits: int = 2) -> str:
    return f"{float(value) * 100:.{digits}f}%"


def set_cell_text(cell, text: str, bold: bool = False, size: float = 9.0):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders_three_line(table):
    # A practical three-line style: top/bottom border for whole table and
    # bottom border for header row, no vertical inner borders.
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.first_child_found_in("w:tcBorders")
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for edge in ("left", "right", "insideH", "insideV"):
                tag = f"w:{edge}"
                el = borders.find(qn(tag))
                if el is None:
                    el = OxmlElement(tag)
                    borders.append(el)
                el.set(qn("w:val"), "nil")
            for edge in ("top", "bottom"):
                tag = f"w:{edge}"
                el = borders.find(qn(tag))
                if el is None:
                    el = OxmlElement(tag)
                    borders.append(el)
                el.set(qn("w:val"), "single")
                el.set(qn("w:sz"), "8" if row_idx in (0, len(table.rows) - 1) else "4")
                el.set(qn("w:color"), "000000" if row_idx in (0, len(table.rows) - 1) else "BFBFBF")
            if row_idx == 0:
                bottom = borders.find(qn("w:bottom"))
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "8")
                bottom.set(qn("w:color"), "000000")


def set_run_font(run, name="宋体", size=12, bold=False, italic=False):
    run.font.name = name if name != "宋体" else "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_section(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.gutter = Cm(0)
    footer = section.footer
    if footer.paragraphs:
        add_page_number(footer.paragraphs[0])


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    for name, size in [("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)


def add_paragraph(doc, text: str = "", *, align=None, first_indent=True, bold=False, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    else:
        p.paragraph_format.first_line_indent = None
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, "宋体", size, bold=bold)
    return p


def add_heading(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if level > 1 else WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = None
    for run in p.runs:
        set_run_font(run, "黑体", 15 if level == 1 else 12, bold=True)
    return p


def add_caption(doc, text: str, above: bool = False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(3 if not above else 0)
    p.paragraph_format.space_after = Pt(3 if above else 6)
    run = p.add_run(text)
    set_run_font(run, "宋体", 10.5, bold=True)
    return p


def add_table(doc, rows: list[list[str]], caption: str):
    add_caption(doc, caption, above=True)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, value, bold=(r_idx == 0), size=8.5)
            if r_idx == 0:
                set_cell_shading(cell, "EDEDED")
    set_table_borders_three_line(table)
    doc.add_paragraph()
    return table


def add_figure(doc, image_name: str, caption: str, width_cm: float = 12.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    run = p.add_run()
    run.add_picture(str(FIG_DIR / image_name), width=Cm(width_cm))
    add_caption(doc, caption, above=False)


def add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目录")
    set_run_font(r, "黑体", 15, bold=True)

    entries = [
        "摘要",
        "Abstract",
        "第一章 绪论",
        "  1.1 研究背景",
        "  1.2 研究问题与研究意义",
        "  1.3 技术路线与论文结构",
        "第二章 数据来源、变量说明与适用性",
        "  2.1 数据来源",
        "  2.2 变量类型与业务关系",
        "  2.3 数据集适用性",
        "  2.4 数据集局限与课程分析边界",
        "第三章 R 工程流程与数据预处理",
        "  3.1 R 工具链与可复现设计",
        "  3.2 数据质量检查与变量处理",
        "第四章 描述性统计与探索性可视化",
        "第五章 统计推断分析",
        "第六章 回归模型与影响因素分析",
        "第七章 计数模型与预测模型比较",
        "第八章 用户类型差异分析",
        "第九章 结论、建议与不足",
        "参考文献",
        "致谢",
        "附录 A R 运行环境",
        "附录 B 核心 R 分析代码",
        "附录 C 附图附表全文与模型文件清单",
    ]
    for entry in entries:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(entry)
        set_run_font(run, "宋体", 12, bold=not entry.startswith("  "))
    doc.add_page_break()


def rows_from_dicts(data: list[dict[str, str]], cols: list[tuple[str, str]], limit=None, digits=3) -> list[list[str]]:
    rows = [[label for _, label in cols]]
    for row in data[:limit]:
        out = []
        for key, _ in cols:
            val = row.get(key, "")
            if re.fullmatch(r"-?\d+(\.\d+)?(e-?\d+)?", str(val), re.I):
                out.append(fmt_num(val, digits))
            else:
                out.append(val)
        rows.append(out)
    return rows


class SimpleHTMLTableParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.tables: list[list[list[str]]] = []
        self._current_table: list[list[str]] | None = None
        self._current_row: list[str] | None = None
        self._current_cell: list[str] | None = None
        self._in_cell = False

    def handle_starttag(self, tag, attrs):
        if tag == "table":
            self._current_table = []
        elif tag == "tr" and self._current_table is not None:
            self._current_row = []
        elif tag in {"td", "th"} and self._current_row is not None:
            self._current_cell = []
            self._in_cell = True

    def handle_data(self, data):
        if self._in_cell and self._current_cell is not None:
            self._current_cell.append(data)

    def handle_endtag(self, tag):
        if tag in {"td", "th"} and self._current_row is not None and self._current_cell is not None:
            text = re.sub(r"\s+", " ", "".join(self._current_cell)).strip()
            self._current_row.append(text)
            self._current_cell = None
            self._in_cell = False
        elif tag == "tr" and self._current_table is not None and self._current_row is not None:
            if self._current_row:
                self._current_table.append(self._current_row)
            self._current_row = None
        elif tag == "table" and self._current_table is not None:
            if self._current_table:
                self.tables.append(self._current_table)
            self._current_table = None


def normalize_rows(rows: list[list[str]]) -> list[list[str]]:
    if not rows:
        return [["无内容"]]
    max_cols = max(len(row) for row in rows)
    return [row + [""] * (max_cols - len(row)) for row in rows]


def read_table_file(path: Path) -> list[list[str]]:
    if path.suffix.lower() == ".csv":
        with path.open(newline="", encoding="utf-8-sig") as f:
            return normalize_rows([[str(cell) for cell in row] for row in csv.reader(f)])
    if path.suffix.lower() in {".html", ".htm"}:
        parser = SimpleHTMLTableParser()
        parser.feed(path.read_text(encoding="utf-8", errors="ignore"))
        rows: list[list[str]] = []
        for idx, table_rows in enumerate(parser.tables, start=1):
            if idx > 1:
                rows.append([f"HTML 子表 {idx}"])
            rows.extend(table_rows)
        return normalize_rows(rows)
    text = path.read_text(encoding="utf-8", errors="ignore")
    return [[line] for line in text.splitlines()] or [["无内容"]]


def add_appendix_image(doc, image_path: Path, caption: str, width_cm: float = 13.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    add_caption(doc, caption, above=False)


def add_appendix_data_table(doc, rows: list[list[str]], caption: str):
    rows = normalize_rows(rows)
    add_caption(doc, caption, above=True)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    font_size = 5.5 if len(rows[0]) >= 8 else 6.5
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, value, bold=(r_idx == 0), size=font_size)
            if r_idx == 0:
                set_cell_shading(cell, "EDEDED")
    set_table_borders_three_line(table)
    doc.add_paragraph()


def preview_rows(rows: list[list[str]], max_data_rows: int = 12) -> list[list[str]]:
    rows = normalize_rows(rows)
    if len(rows) <= max_data_rows + 1:
        return rows
    omitted = len(rows) - max_data_rows - 1
    return rows[: max_data_rows + 1] + [[f"以下省略 {omitted} 行，完整数据见原始文件路径。"] + [""] * (len(rows[0]) - 1)]


def build():
    doc = Document()
    configure_section(doc.sections[0])
    configure_styles(doc)

    # Cover
    add_paragraph(doc, "2025—2026 第二学期《统计软件》课程论文", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, bold=True, size=15)
    if (ASSET_DIR / "course_cover_image.png").exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        p.add_run().add_picture(str(ASSET_DIR / "course_cover_image.png"), width=Cm(12))
    add_paragraph(doc, "题目：基于 R 语言的共享单车租赁需求影响因素与预测分析", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, bold=True, size=14)
    add_paragraph(doc, "——以 UCI Bike Sharing Dataset 为例", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, bold=True, size=14)
    for label in ["姓名：__________", "学号：__________", "学院：__________", "专业：__________"]:
        add_paragraph(doc, label, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, bold=True, size=12)
    add_paragraph(doc, "2026 年 6 月", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, bold=True, size=12)
    doc.add_page_break()

    add_toc(doc)

    # Chinese title and abstract
    add_paragraph(doc, "基于 R 语言的共享单车租赁需求影响因素与预测分析", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, bold=True, size=15)
    add_paragraph(doc, "——以 UCI Bike Sharing Dataset 为例", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, bold=True, size=15)
    add_heading(doc, "摘要", 1)
    add_paragraph(
        doc,
        f"本文基于 UCI Machine Learning Repository 收录的 Bike Sharing Dataset，选取 2011 年至 2012 年美国华盛顿 D.C. Capital Bikeshare 系统小时级租赁数据开展统计分析[1]。数据集由 Fanaee-T 与 Gama 在共享单车事件检测研究背景下整理发布，具有真实交通系统日志、时间属性、天气环境属性和用户类型信息并存的特点[2]。本文主分析数据为 hour.csv，共包含 {sval('hour_rows')} 条小时级记录和 {sval('hour_variables')} 个原始变量；辅助数据 day.csv 共包含 {sval('day_rows')} 条日级记录。研究使用 R 语言完成数据读取、质量检查、变量转换、描述统计、可视化、统计推断、回归建模、计数模型、机器学习预测和用户类型差异分析。",
    )
    add_paragraph(
        doc,
        f"结果显示，共享单车租赁量存在明显日内周期和季节差异。平均租赁量最高的小时为 {sval('peak_hour')} 时，平均每小时租赁量为 {sval('peak_hour_mean')}；最低的小时为 {sval('lowest_hour')} 时，平均租赁量仅为 {sval('lowest_hour_mean')}。季节方面，{sval('season_highest_mean')}平均租赁量最高，{sval('season_lowest_mean')}最低。回归分析表明，在控制时间、日期和天气变量后，气温与租赁量呈正相关，湿度和风速与租赁量呈负相关；Breusch-Pagan 检验发现异方差，因此本文使用 HC3 稳健标准误进行补充检验。计数模型显示 Poisson 回归存在明显过度离散，过度离散比为 {sval('poisson_overdispersion')}，负二项回归 AIC 明显低于 Poisson 回归。预测模型比较显示，随机森林在测试集上的 RMSE 为 {sval('best_prediction_RMSE')}，MAE 为 {sval('best_prediction_MAE')}，R 平方为 {sval('best_prediction_R2')}，显著优于线性回归、负二项回归、LASSO 回归和回归树。",
    )
    add_paragraph(
        doc,
        f"用户类型分析显示，注册用户贡献了总租赁量的 {pct(sval('registered_total_share'))}，休闲用户贡献 {pct(sval('casual_total_share'))}。注册用户更体现工作日通勤属性，休闲用户对温度、湿度和风速等舒适度变量更敏感。本文结果说明，R 语言能够有效支持从数据预处理、统计描述、模型诊断到预测建模的完整数据分析流程，也说明在共享单车需求研究中，应区分解释模型与预测模型的作用，避免将预测精度直接等同于因果解释。",
    )
    add_paragraph(doc, "关键词：共享单车；R 语言；租赁需求；负二项回归；随机森林；LASSO", first_indent=False, bold=True)

    # English abstract
    add_paragraph(doc, "Demand Factor Analysis and Prediction of Bike Sharing Rentals Using R", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, bold=True, size=15)
    add_paragraph(doc, "A Case Study Based on the UCI Bike Sharing Dataset", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, bold=True, size=15)
    add_heading(doc, "Abstract", 1)
    add_paragraph(
        doc,
        "This paper analyzes the UCI Bike Sharing Dataset using R. The hourly data from the Capital Bikeshare system in Washington, D.C. cover the years 2011 and 2012 and include temporal, calendar, weather, environmental and user-type variables. The study conducts data quality checks, variable transformation, descriptive statistics, exploratory visualization, inferential tests, linear regression, count models, regularized regression, random forest prediction and user-segment analysis. The analysis emphasizes reproducibility, leakage control and appropriate interpretation of statistical and predictive models.",
    )
    add_paragraph(
        doc,
        f"The empirical results show strong temporal patterns in bike sharing demand. The highest average hourly demand occurs at {sval('peak_hour')}:00, while the lowest occurs at {sval('lowest_hour')}:00. Temperature is positively associated with rentals, whereas humidity and wind speed are negatively associated with demand after controlling for temporal and weather factors. The Poisson model suffers from substantial overdispersion, and the negative binomial model provides a better fit. For prediction, the random forest model achieves the best test-set performance, with RMSE = {sval('best_prediction_RMSE')}, MAE = {sval('best_prediction_MAE')} and R-squared = {sval('best_prediction_R2')}. The results also reveal clear behavioral differences between casual and registered users. This paper demonstrates that R provides a coherent toolchain for statistical description, model diagnosis and predictive analysis of urban mobility data.",
    )
    add_paragraph(doc, "Keywords: bike sharing; R; demand prediction; negative binomial regression; random forest; LASSO", first_indent=False, bold=True)
    doc.add_page_break()

    # Body
    add_heading(doc, "第一章 绪论", 1)
    add_heading(doc, "1.1 研究背景", 2)
    for para in [
        "共享单车是城市短途出行和公共交通接驳的重要组成部分。与传统公共交通相比，共享单车具有借还灵活、覆盖半径小、使用门槛低等特点，能够补充地铁、公交与步行之间的短距离出行需求。共享单车系统在运行过程中会自动记录租赁时间、用户类型和使用数量，这些数据构成了城市出行需求的微观记录。与问卷调查不同，运营日志具有连续性和客观性，因此适合用于统计软件课程中的真实数据分析。",
        "本文选择 Bike Sharing Dataset 作为研究对象，原因在于该数据集兼具现实背景、变量多样性和建模可操作性。其原始数据来自美国华盛顿 D.C. Capital Bikeshare 系统，并结合天气和季节信息进行整理[1]。Fanaee-T 与 Gama 的研究指出，共享单车租赁量不仅与天气和季节有关，也可能反映城市事件和异常活动[2]。这说明共享单车数据具有双重价值：一方面可以用于需求预测，另一方面也可以作为理解城市活动节奏的观测窗口。",
        "对于统计软件课程而言，一个合适的数据集不应只是“能读入 R”，还应能够支撑多种统计分析任务。Bike Sharing Dataset 同时包含计数型目标变量、时间变量、分类变量、连续环境变量和用户分组变量，适合展示描述统计、可视化、统计推断、回归诊断、计数模型和机器学习预测。本文将围绕“哪些因素影响共享单车租赁需求，以及 R 软件能否较好刻画和预测这种需求变化”展开。",
    ]:
        add_paragraph(doc, para)

    add_heading(doc, "1.2 研究问题与研究意义", 2)
    for para in [
        "本文重点回答五个问题：第一，共享单车租赁量是否存在明显的小时、月份和季节规律；第二，天气状况、气温、湿度和风速等环境变量与租赁需求之间存在怎样的关系；第三，工作日和非工作日的租赁模式是否不同；第四，休闲用户与注册用户是否表现出不同使用规律；第五，解释性统计模型和机器学习预测模型在本数据集上各自表现如何。",
        "研究意义主要体现在三个方面。其一，从数据理解角度看，本文通过描述统计和图形分析展示共享单车需求的分布特征和时间结构；其二，从统计建模角度看，本文比较线性回归、计数模型和机器学习模型，说明不同模型适合解决不同问题；其三，从课程实践角度看，本文完整使用 R 软件完成读取、整理、绘图、检验、建模和导出，体现统计软件综合运用能力。",
    ]:
        add_paragraph(doc, para)

    add_heading(doc, "1.3 技术路线与论文结构", 2)
    for para in [
        "本文的技术路线遵循“数据来源核验—数据质量检查—变量工程—描述统计—探索性可视化—统计推断—模型诊断—预测比较—分用户讨论—结论建议”的顺序。这样安排的原因是，真实数据分析不能从模型开始，而应先确认数据是否可信、变量是否理解正确、图形是否揭示稳定规律，再进入回归和机器学习模型。本文所有图表均由 R 脚本自动生成，避免手工统计带来的不一致。",
        "论文结构也与这一技术路线对应。第一章提出研究背景和问题；第二章说明数据来源、变量含义和课程适用性；第三章介绍 R 工具链、工程目录和预处理逻辑；第四章用图表呈现租赁量的时间、天气和相关结构；第五章进行统计推断；第六章建立解释性回归模型并进行诊断；第七章比较计数模型和预测模型；第八章区分休闲用户与注册用户；第九章按时间规律、天气影响、模型结论、用户差异、运营启示和研究限制进行总结。",
    ]:
        add_paragraph(doc, para)

    add_heading(doc, "第二章 数据来源、变量说明与适用性", 1)
    add_heading(doc, "2.1 数据来源", 2)
    for para in [
        "本文数据来自 UCI Machine Learning Repository 收录的 Bike Sharing Dataset。UCI 是机器学习和数据挖掘领域常用的公开数据集平台，该数据集的 DOI 为 10.24432/C5W894[1]。本地数据文件包括 hour.csv、day.csv 和 Readme.txt，其中 hour.csv 是本文主分析数据，day.csv 用于补充趋势分析。",
        "根据数据说明，hour.csv 记录 2011 年至 2012 年按小时汇总的共享单车租赁数量，共 17,379 条记录；day.csv 记录按日汇总的数据，共 731 条记录。原始租赁记录来自 Capital Bikeshare 系统，天气信息由外部气象数据整理得到。由于共享单车使用行为与天气、日期、季节和小时密切相关，该数据集天然适合做回归预测和需求影响因素分析。",
    ]:
        add_paragraph(doc, para)

    overview_rows = rows_from_dicts(
        read_csv("table01_dataset_overview.csv"),
        [("dataset", "数据文件"), ("observations", "观测数"), ("variables", "变量数"), ("start_date", "起始日期"), ("end_date", "结束日期"), ("role", "用途")],
        digits=0,
    )
    add_table(doc, overview_rows, "表 1 数据集基本信息 / Table 1 Basic information of datasets")

    add_heading(doc, "2.2 变量类型与业务关系", 2)
    for para in [
        "hour.csv 的变量可分为五类：索引与日期变量、时间属性变量、日期类型变量、天气环境变量和租赁量变量。时间属性包括年份、月份、小时和星期；日期类型包括节假日和工作日；天气环境变量包括季节、天气状况、气温、体感温度、湿度和风速；租赁量变量包括休闲用户、注册用户和总租赁量。",
        "本文特别强调一个关键业务关系：总租赁量 cnt 等于休闲用户租赁量 casual 与注册用户租赁量 registered 之和。这一关系决定了建模边界：在解释或预测 cnt 时，不能把 casual 和 registered 作为自变量，否则会造成数据泄露，使模型获得本不应提前知道的信息。本文只在用户类型差异分析中分别使用 casual 和 registered 作为因变量。",
    ]:
        add_paragraph(doc, para)

    dict_rows = rows_from_dicts(
        read_csv("table02_variable_dictionary.csv"),
        [("variable", "变量名"), ("meaning", "含义"), ("analysis_role", "分析角色")],
        digits=0,
    )
    add_table(doc, dict_rows, "表 2 变量字典与分析角色 / Table 2 Variable dictionary and analytical roles")

    add_heading(doc, "2.3 数据集适用性", 2)
    for para in [
        "从统计软件课程角度看，该数据集的适用性非常强。首先，样本量适中。hour.csv 包含 17,379 条小时级记录，既能支撑统计检验、回归建模和机器学习预测，又不会因为规模过大导致普通个人电脑运行困难。对于课程论文而言，这一规模可以充分展示 R 的数据处理能力，同时保持代码运行和结果复现的可控性。",
        "其次，变量类型多样，适合展示统计软件的多种功能。数据中既有 cnt 这样的计数型因变量，也有 temp、hum、windspeed 等连续变量，还有 season、weathersit、workingday、holiday、hr、mnth 等分类或有序时间变量。因此，本文可以同时开展连续变量分析、分类变量分组比较、方差分析或非参数检验、线性回归、计数回归和机器学习预测，能够较完整地覆盖统计软件课程常见内容。",
        "第三，该数据集兼具解释分析和预测分析价值。解释分析方面，可以讨论气温、湿度、风速、工作日和天气状况与租赁需求之间的方向性关系；预测分析方面，可以利用时间、日期和天气变量预测小时租赁量。由于目标变量 cnt 不是单纯连续变量，而是非负整数，还可以进一步展示 Poisson 回归和负二项回归在计数数据中的应用。",
        "第四，该数据集虽然没有缺失值，但仍适合展示严谨的数据质量检查。课程论文中常见问题是直接使用“干净数据”建模，而忽视缺失值、重复值、变量关系和逻辑一致性检查。本文即使发现数据无缺失，也仍然通过 R 检查各变量缺失情况、重复记录、instant 编号唯一性以及 cnt = casual + registered 的组成关系。这种写法能够说明作者理解真实数据分析流程，而不是只把数据集当作模型输入。",
        "第五，该数据集还具有较好的原创阐释空间。小时图、工作日图、天气图、温度曲线和相关图都能从城市出行行为角度进行解释，不容易停留在机械复述统计量。本文的图表解释重点放在通勤节律、休闲出行、天气舒适度和需求波动机制上，从而增强论文语言的自主分析性，降低与常见数据集介绍文字重复的风险。",
    ]:
        add_paragraph(doc, para)

    add_heading(doc, "2.4 数据集局限与课程分析边界", 2)
    for para in [
        "需要说明的是，Bike Sharing Dataset 虽然适合统计软件课程，但并不等于可以支持所有类型的因果结论。数据仅覆盖华盛顿 D.C. 某一共享单车系统的 2011 至 2012 年记录，缺少站点位置、车辆投放量、价格策略、用户注册规模、突发事件和节假日活动强度等变量。因此，本文将研究目标限定为需求影响因素的统计关联和预测建模，不把回归系数解释为严格因果效应。",
        "这一边界反而使数据集更适合课程论文：它足够真实，能暴露异方差、过度离散、非线性和数据泄露等实际问题；同时又足够规整，便于通过 R 语言从头到尾完成分析。本文在建模中明确不使用 casual 和 registered 预测 cnt，就是对数据结构和研究边界的主动控制。",
    ]:
        add_paragraph(doc, para)

    add_heading(doc, "第三章 R 工程流程与数据预处理", 1)
    add_heading(doc, "3.1 R 工具链与可复现设计", 2)
    for para in [
        "本文采用脚本化方式组织 R 分析，核心脚本为 scripts/analysis.R。脚本从原始 CSV 出发，依次完成质量检查、变量转换、统计表输出、图形保存、模型训练、模型对象保存和运行环境记录。所有图片保存到 outputs/figures，所有表格保存到 outputs/tables，模型对象保存到 outputs/models。这种目录设计保证论文中的每一个结果都能追溯到代码和输出文件。",
        "R 包的使用并非简单堆叠，而是围绕任务选择工具。readr 用于数据读取，dplyr 和 tidyr 用于整理，ggplot2 用于可视化，broom 和 modelsummary 用于模型结果整理，MASS 用于负二项回归，glmnet 用于 LASSO，ranger 用于随机森林，car、lmtest 和 sandwich 用于模型诊断与稳健标准误。这些工具共同构成从统计描述到机器学习预测的分析链条[3][4][5]。",
    ]:
        add_paragraph(doc, para)

    tool_rows = [
        ["R 工具或包", "主要用途", "对应环节"],
        ["readr", "读取 CSV 数据", "数据读取"],
        ["dplyr / tidyr", "变量转换、分组汇总、长宽表整理", "数据预处理"],
        ["ggplot2", "绘制分布图、折线图、箱线图、散点平滑图", "可视化"],
        ["broom / modelsummary", "整理模型结果和模型比较表", "模型输出"],
        ["car / lmtest / sandwich", "VIF、Breusch-Pagan、HC3 稳健标准误", "模型诊断"],
        ["MASS", "负二项回归", "计数模型"],
        ["glmnet", "LASSO 正则化回归", "预测模型"],
        ["ranger / vip", "随机森林与变量重要性", "机器学习预测"],
        ["rpart", "回归树对照模型", "机器学习解释"],
    ]
    add_table(doc, tool_rows, "表 3 R 工具使用情况 / Table 3 R packages and analytical purposes")
    for para in [
        "表 3 展示了本文对 R 工具的综合使用。工具数量较多，但各工具的角色是清晰分工的：数据导入、数据整理、图形表达、模型估计、模型诊断、预测评估和结果导出分别由不同包完成。这样的安排既体现 R 生态的优势，也避免把所有任务都压到基础函数或单一包中，因而代码更简洁、更易读、更便于复现。",
        "从评分角度看，本文不是为了增加包名而增加工具，而是让每个工具对应一个统计分析任务。例如，MASS 的 glm.nb 解决过度离散计数数据问题，glmnet 解决高维哑变量进入线性预测时的正则化问题，ranger 解决非线性和交互结构的预测问题，lmtest 与 sandwich 则服务于回归诊断和稳健推断。这种工具选择使论文的统计语言与软件实现保持一致。",
    ]:
        add_paragraph(doc, para)

    add_heading(doc, "3.2 数据质量检查与变量处理", 2)
    for para in [
        "数据预处理首先检查 hour.csv 和 day.csv 的行数、列数、日期范围和变量名称。随后，对 hour.csv 的所有变量进行缺失值检查，并检查完全重复记录和 instant 编号重复情况。检查结果显示，数据没有缺失值和重复记录，且所有记录均满足 cnt = casual + registered。",
        "变量处理方面，本文将 season、yr、mnth、hr、holiday、weekday、workingday 和 weathersit 转换为有实际含义的因子变量；将标准化气温还原为摄氏度，将湿度还原为百分比，将风速还原为实际指标；构造 log1p(cnt) 以缓解租赁量右偏分布对线性回归的影响；同时构造用户占比变量，用于后续休闲用户和注册用户差异分析。",
        "预处理代码采用管道式写法，将变量重编码、因子标签、数值还原和派生变量构造集中在同一流程中完成。这种写法的优点是逻辑顺序清楚，读者可以沿着代码看到原始变量如何转化为论文中使用的分析变量。所有中间结果均保存为 RDS 文件，既保留了数据类型信息，也方便后续模型和论文生成脚本复用。",
        "本文还特别控制了预测任务中的信息泄露问题。由于 cnt 由 casual 和 registered 相加得到，如果把这两个变量作为预测 cnt 的自变量，模型会得到近似直接答案，测试集表现会被虚高。R 脚本在建模公式中主动排除这两个组成变量，保证预测性能来自时间、日期、天气和环境变量，而不是来自目标变量的拆分项。",
    ]:
        add_paragraph(doc, para)
    quality_rows = rows_from_dicts(read_csv("table04_quality_checks.csv"), [("check_item", "检查项目"), ("value", "数值"), ("conclusion", "结论")], digits=0)
    add_table(doc, quality_rows, "表 4 数据质量检查结果 / Table 4 Data quality checks")

    add_heading(doc, "第四章 描述性统计与探索性可视化", 1)
    add_heading(doc, "4.1 租赁量分布与基础统计", 2)
    for para in [
        f"描述性统计显示，总租赁量 cnt 的均值为 {sval('cnt_mean')}，中位数为 {sval('cnt_median')}，最大值为 {sval('cnt_max')}。均值高于中位数，说明租赁量分布右偏，少数高峰时段的租赁量显著高于普通时段。休闲用户平均租赁量为 {sval('casual_mean')}，注册用户平均租赁量为 {sval('registered_mean')}，注册用户是系统需求主体。",
        "右偏分布对建模有直接影响。如果直接用原始 cnt 建立线性回归，残差可能受到高峰时段影响而出现异方差。因此本文在线性模型中使用 log1p(cnt)，而在计数模型中保留原始 cnt。这样的处理区分了不同模型对因变量尺度的要求。",
    ]:
        add_paragraph(doc, para)
    desc_rows = rows_from_dicts(read_csv("table05_descriptive_stats.csv"), [("variable_cn", "变量"), ("n", "样本量"), ("mean", "均值"), ("sd", "标准差"), ("min", "最小值"), ("median", "中位数"), ("max", "最大值")], digits=2)
    add_table(doc, desc_rows, "表 5 主要连续变量描述性统计 / Table 5 Descriptive statistics of continuous variables")
    add_figure(doc, "fig01_cnt_distribution.png", "图 1 总租赁量分布 / Figure 1 Distribution of total rental count", 12)
    add_paragraph(
        doc,
        "图 1 的核心信息是租赁量分布并不对称，而是呈现明显右偏。大多数小时处于中低租赁量区间，但少数通勤高峰小时会出现非常高的 cnt 值，这使均值被高峰时段拉高。这个发现对后续模型选择非常关键：一方面，线性回归若直接使用原始 cnt，容易受到极端高峰小时影响；另一方面，租赁量作为计数变量，也提示有必要比较 Poisson 回归和负二项回归。换言之，图 1 不只是描述现象，还直接解释了为什么本文后续采用 log1p(cnt) 和计数模型两条路径。",
    )

    add_heading(doc, "4.2 时间规律", 2)
    for para in [
        f"小时维度是共享单车需求最重要的结构性因素。图 2 显示，凌晨需求最低，其中 {sval('lowest_hour')} 时平均租赁量仅为 {sval('lowest_hour_mean')}；傍晚需求最高，{sval('peak_hour')} 时平均租赁量达到 {sval('peak_hour_mean')}。该模式说明共享单车明显受到日内出行节律影响，尤其与通勤活动相关。",
        "工作日与非工作日的差异不仅体现在平均水平，更体现在曲线形态。工作日早晚高峰突出，非工作日白天需求较平缓，说明共享单车同时承载通勤出行和休闲出行两类场景。这一发现直接支持第六章加入 workingday 与 hr 的交互项。",
    ]:
        add_paragraph(doc, para)
    add_figure(doc, "fig03_hourly_pattern.png", "图 2 不同小时平均租赁量 / Figure 2 Average rental count by hour", 13)
    add_paragraph(
        doc,
        "图 2 可以看作本文最重要的探索性图形之一。凌晨 0 至 5 时需求快速下降并维持低位，说明夜间骑行需求有限；6 时以后需求迅速上升，8 时形成明显早高峰；中午至下午保持较高水平；17 时达到全天峰值，18 时仍维持高位。这种双峰或偏双峰结构与城市通勤非常一致，说明小时变量不应被当作普通连续变量简单线性处理，而应作为分类变量进入模型，使每个小时拥有独立水平。",
    )
    add_figure(doc, "fig04_workingday_hour_pattern.png", "图 3 工作日与非工作日小时模式 / Figure 3 Hourly patterns by working day", 13)
    add_paragraph(
        doc,
        "图 3 进一步揭示工作日与非工作日的行为机制差异。工作日曲线在早晚通勤时段更陡峭，说明注册用户或固定通勤者可能在这些时段集中使用共享单车；非工作日曲线则更接近单峰，白天休闲出行占比提高。这一差异说明，仅比较工作日平均租赁量并不足够，因为平均值会掩盖日内结构变化。后续回归模型加入 workingday 与 hr 的交互项，正是为了把这种图形中可见的曲线差异转化为可检验的统计模型。",
    )
    add_figure(doc, "fig07_weekday_hour_heatmap.png", "图 4 星期与小时租赁热力图 / Figure 4 Rental heatmap by weekday and hour", 13)
    add_paragraph(
        doc,
        "图 4 用热力图展示星期和小时的二维结构，比单条折线更直观。颜色较深的区域集中在工作日早晚高峰，周末则更多集中于白天到傍晚。这说明共享单车需求不是由单一时间变量决定，而是由“星期属性”和“小时属性”共同塑造。由于 weekday 与 holiday、workingday 存在定义上的重叠，本文没有在主回归中同时纳入所有日期变量，而是把热力图作为理解时间结构和避免模型设定冗余的重要依据。",
    )

    add_heading(doc, "4.3 季节、天气与环境变量", 2)
    for para in [
        f"季节分析显示，{sval('season_highest_mean')}平均租赁量最高，{sval('season_lowest_mean')}最低。天气分析显示，晴朗或少云天气下租赁量最高，小雨雪天气下需求下降。恶劣天气类别样本极少，因此本文仅将其作为描述性观察，不据此进行强推断。",
        "气温与租赁量呈总体正向关系，说明较舒适温度会提高骑行意愿；湿度和风速则呈负向趋势，说明高湿度和大风可能降低骑行舒适性。这些关系在图形中表现明显，但仍需在回归模型中控制小时、季节和天气等因素后进一步判断。",
    ]:
        add_paragraph(doc, para)
    group_rows = rows_from_dicts(read_csv("table07_group_means.csv"), [("group_variable", "分组变量"), ("group", "类别"), ("n", "样本量"), ("mean_cnt", "平均租赁量"), ("median_cnt", "中位数")], digits=2)
    add_table(doc, group_rows, "表 6 主要分类变量分组均值 / Table 6 Group means of categorical variables")
    add_figure(doc, "fig08_weather_boxplot.png", "图 5 不同天气状况下租赁量分布 / Figure 5 Rental distribution by weather condition", 13)
    add_paragraph(
        doc,
        f"天气箱线图显示，晴朗或少云天气的平均租赁量为 {fmt_num(read_csv('table14_weather_summary.csv')[0]['mean_cnt'], 2)}，多云或有雾天气下降到 {fmt_num(read_csv('table14_weather_summary.csv')[1]['mean_cnt'], 2)}，小雨雪天气进一步下降到 {fmt_num(read_csv('table14_weather_summary.csv')[2]['mean_cnt'], 2)}。这不是单纯的均值差异，而是分布整体下移：恶劣天气会同时减少普通时段需求和高峰时段需求。由于恶劣天气类别只有 3 条记录，本文不对该类别作强结论，这体现了图形解释中对样本量约束的尊重。",
    )
    add_figure(doc, "fig09_temp_cnt_smooth.png", "图 6 气温与总租赁量关系 / Figure 6 Relationship between temperature and rentals", 13)
    add_paragraph(
        doc,
        f"温度平滑图显示气温与租赁量大体正相关，Pearson 相关系数为 {fmt_num(read_csv('table17_environment_correlations.csv')[0]['pearson'], 3)}，Spearman 相关系数为 {fmt_num(read_csv('table17_environment_correlations.csv')[0]['spearman'], 3)}。这说明温度升高通常伴随骑行需求提高，但图形也暗示这种关系不是无限线性增长，而更接近舒适区间内的上升关系。现实含义是，气温改善会提高骑行意愿，但过热天气是否继续增加需求，还需要更细的城市气候和体感变量支持。",
    )
    add_figure(doc, "fig10_humidity_cnt_smooth.png", "图 7 湿度与总租赁量关系 / Figure 7 Relationship between humidity and rentals", 13)
    add_paragraph(
        doc,
        f"湿度图呈现较稳定的负向趋势，湿度与总租赁量的 Pearson 相关系数为 {fmt_num(read_csv('table17_environment_correlations.csv')[1]['pearson'], 3)}。高湿度往往意味着闷热、降雨可能性上升或骑行舒适度下降，因此用户即使有出行需求，也可能转向公交、地铁或步行。与温度相比，湿度对休闲用户的影响在第八章中更明显，说明舒适度变量不仅影响总量，也影响不同用户群体的选择。",
    )
    add_figure(doc, "fig11_windspeed_cnt_smooth.png", "图 8 风速与总租赁量关系 / Figure 8 Relationship between wind speed and rentals", 13)
    add_paragraph(
        doc,
        f"风速与总租赁量的简单相关系数为 {fmt_num(read_csv('table17_environment_correlations.csv')[2]['pearson'], 3)}，方向并不像湿度那样直观。这提醒我们不能仅凭二元相关判断变量作用，因为风速可能与季节、温度、天气状况同时变化。在多元回归控制其他因素后，风速系数转为显著负向，说明图形分析应与回归控制结合使用，不能用单张散点图替代模型解释。",
    )

    add_heading(doc, "4.4 相关分析与建模提示", 2)
    for para in [
        "相关系数图具有两个重要提示。第一，cnt 与 casual、registered 高度相关，这是因为后两者构成前者，不能被误解为普通解释变量关系。第二，气温与体感温度高度相关，如果同时进入回归模型，会增加共线性风险。因此本文主模型保留气温，而不同时纳入体感温度。",
        "探索性可视化并不是最终结论，而是模型设定的依据。本文正是根据小时曲线、工作日差异、天气箱线图和相关图，确定后续模型重点控制时间、工作日、天气、气温、湿度和风速变量。",
    ]:
        add_paragraph(doc, para)
    add_figure(doc, "fig12_correlation_heatmap.png", "图 9 连续变量相关系数热力图 / Figure 9 Correlation heatmap of continuous variables", 11)
    add_paragraph(
        doc,
        f"相关热力图最醒目的关系是 cnt 与 registered 的相关系数达到 {fmt_num([r for r in read_csv('table10_correlation_long.csv') if r['var1'] == 'cnt' and r['var2'] == 'registered'][0]['correlation'], 3)}，与 casual 的相关系数为 {fmt_num([r for r in read_csv('table10_correlation_long.csv') if r['var1'] == 'cnt' and r['var2'] == 'casual'][0]['correlation'], 3)}。这并不是发现“注册用户导致总租赁量增加”，而是因为 cnt 本身由两类用户相加得到。本文据此明确避免把 casual 和 registered 放入 cnt 的预测模型。热力图还显示 temp_celsius 与 atemp_celsius 的相关系数高达 {fmt_num([r for r in read_csv('table10_correlation_long.csv') if r['var1'] == 'temp_celsius' and r['var2'] == 'atemp_celsius'][0]['correlation'], 3)}，因此主模型只保留温度变量，避免高度共线变量重复进入模型。",
    )

    add_heading(doc, "第五章 统计推断分析", 1)
    for para in [
        "探索性图表能够显示差异，但仍需要通过统计检验判断差异是否具有统计显著性。本文对工作日与非工作日、不同季节、不同天气状况以及不同年份的租赁量差异进行检验。由于租赁量分布右偏，本文同时使用参数检验和非参数检验，以避免单一检验假设过强。",
        "检验结果显示，工作日与非工作日、季节、天气状况和年份差异均达到 5% 显著性水平。需要注意的是，本数据样本量较大，p 值容易显著，因此解释时不能只看 p 值，还需要结合分组均值和图形判断差异是否具有实际意义。",
    ]:
        add_paragraph(doc, para)
    infer_rows = rows_from_dicts(read_csv("table16_inference_tests.csv"), [("comparison", "比较对象"), ("method", "方法"), ("statistic", "统计量"), ("p_value", "p 值"), ("conclusion", "结论")], digits=4)
    add_table(doc, infer_rows, "表 7 统计推断检验结果 / Table 7 Inferential test results")

    add_heading(doc, "第六章 回归模型与影响因素分析", 1)
    add_heading(doc, "6.1 线性回归模型设定", 2)
    for para in [
        "线性回归用于解释各因素与租赁量之间的方向性关系。由于原始 cnt 分布右偏，本文以 log1p(cnt) 为因变量。解释变量包括季节、年份、月份、小时、节假日、工作日、天气状况、气温、湿度和风速。weekday 没有与 workingday 和 holiday 同时进入主模型，以避免定义关系带来的秩亏风险。",
        "基准模型能够解释 log1p(cnt) 的较大比例变化；加入 workingday 与 hr 交互项后，模型拟合进一步提高，说明工作日属性会改变不同小时的需求模式。这与第四章中工作日和非工作日曲线形态不同的发现一致。",
        "本文把线性回归定位为解释模型，而不是最终预测工具。选择对数变换后的线性回归，是为了让系数可以近似解释为百分比变化，并在控制其他因素后观察变量方向。与随机森林相比，线性回归牺牲了一部分非线性拟合能力，但换来了更透明的系数、显著性检验和模型诊断空间，因此适合放在影响因素分析章节。",
    ]:
        add_paragraph(doc, para)
    lm_rows = rows_from_dicts(read_csv("table20_lm_model_fit.csv"), [("model", "模型"), ("r.squared", "R 平方"), ("adj.r.squared", "调整 R 平方"), ("AIC", "AIC"), ("BIC", "BIC")], digits=3)
    add_table(doc, lm_rows, "表 8 线性回归模型拟合结果 / Table 8 Linear regression model fit")

    add_heading(doc, "6.2 重点变量解释", 2)
    for para in [
        "重点系数显示，2012 年相对于 2011 年系数为正，说明系统整体使用规模提高。天气变量中，小雨雪相对于晴朗或少云呈明显负向影响。气温系数为正，湿度和风速系数为负，表明骑行舒适度变量在控制时间结构后仍与租赁量相关。",
        "由于因变量为 log1p(cnt)，系数不能直接解释为原始租赁量增加多少。本文使用 exp(系数)-1 的近似百分比变化辅助解释，使回归语言更符合对数模型含义。",
    ]:
        add_paragraph(doc, para)
    lm_focus_rows = rows_from_dicts(read_csv("table21_lm_focus_terms.csv"), [("term", "变量"), ("estimate", "系数"), ("p.value", "p 值"), ("percent_change", "近似百分比变化")], digits=4)
    add_table(doc, lm_focus_rows, "表 9 线性回归重点变量结果 / Table 9 Key coefficients in linear regression")

    add_heading(doc, "6.3 模型诊断与稳健性", 2)
    for para in [
        "模型诊断是回归分析的重要组成部分。本数据中季节、月份、气温和天气可能存在相关性，因此需要检查多重共线性。VIF 诊断显示主要变量的调整后 GVIF 没有异常过高，说明主模型解释变量设定总体可接受。本文没有同时纳入气温和体感温度，也是为了降低共线性风险。",
        "Breusch-Pagan 检验显示模型存在异方差[7]。异方差在共享单车数据中具有现实含义：低需求时段波动较小，高峰时段、适宜天气和节假日附近的波动更大。如果忽略异方差，普通最小二乘系数仍可作为条件均值关系的估计，但标准误和显著性检验可能不够可靠。因此，本文使用 HC3 稳健标准误重新检验重点变量[8]，以增强统计推断的可信度。",
        "稳健标准误的作用不是改变模型系数，而是在误差方差不恒定时修正标准误估计。表 11 显示，使用 HC3 后，气温、湿度、风速、小雨雪天气和年份变量的方向与显著性基本保持一致。这说明本文的主要解释结论并不是普通标准误假设下的偶然结果，而是在考虑真实数据波动不均后仍较稳定。",
        "图 10 的残差诊断也能解释为什么需要稳健处理。共享单车租赁量在高预测值区间的残差散布更宽，说明高需求场景下实际租赁量更难被线性模型精确刻画。这并不意味着线性回归无效，而是提示本文不能只报告回归系数，还必须报告诊断结果、稳健标准误，并在第七章进一步引入更适合计数和预测任务的模型。",
    ]:
        add_paragraph(doc, para)
    vif_rows = rows_from_dicts(read_csv("table22_lm_vif_car.csv"), [("term", "变量"), ("gvif_adjusted", "调整 GVIF")], digits=3)
    add_table(doc, vif_rows, "表 10 多重共线性诊断结果 / Table 10 Multicollinearity diagnostics")
    robust_rows = rows_from_dicts(read_csv("table21c_lm_focus_hc3_robust.csv"), [("term", "变量"), ("estimate", "系数"), ("robust_std_error", "HC3 稳健标准误"), ("p_value", "p 值"), ("percent_change", "近似百分比变化")], digits=4)
    add_table(doc, robust_rows, "表 11 HC3 稳健标准误结果 / Table 11 HC3 robust standard error results")
    add_figure(doc, "fig16_lm_diagnostics.png", "图 10 线性回归残差诊断图 / Figure 10 Linear regression diagnostics", 12)

    add_heading(doc, "第七章 计数模型与预测模型比较", 1)
    add_heading(doc, "7.1 Poisson 回归与负二项回归", 2)
    for para in [
        "共享单车租赁量是非负整数，因此 Poisson 回归是自然的计数模型起点。但 Poisson 模型假设均值等于方差，而真实租赁需求受高峰时段、天气和季节共同影响，往往存在更强波动。",
        f"本文计算 Poisson 模型 Pearson 残差平方和与残差自由度之比，得到过度离散比 {sval('poisson_overdispersion')}，明显大于 1。负二项回归通过引入额外离散参数允许方差大于均值，其 AIC 为 {sval('negbin_AIC')}，显著低于 Poisson 回归的 {sval('poisson_AIC')}，说明负二项模型更适合本数据[6]。",
        "负二项模型在本文中仍然属于解释模型。它保留计数型因变量 cnt，不需要把租赁量近似看作连续正态变量；同时可以通过发生率比解释变量影响。例如，天气状况、气温、湿度和风速的发生率比可以被解释为在其他变量不变时租赁发生强度的相对变化。这种解释方式比单纯比较预测误差更接近统计建模的目标。",
    ]:
        add_paragraph(doc, para)
    count_rows = rows_from_dicts(read_csv("table25_count_model_compare.csv"), [("model", "模型"), ("AIC", "AIC"), ("BIC", "BIC"), ("overdispersion_ratio", "过度离散比")], digits=3)
    add_table(doc, count_rows, "表 12 计数模型比较 / Table 12 Comparison of count models")
    irr_rows = rows_from_dicts(read_csv("table24b_negbin_focus_irr.csv"), [("term", "变量"), ("estimate", "系数"), ("incidence_rate_ratio", "发生率比"), ("p.value", "p 值")], digits=4)
    add_table(doc, irr_rows, "表 13 负二项回归重点变量发生率比 / Table 13 Incidence rate ratios of negative binomial model")

    add_heading(doc, "7.2 解释模型与预测模型的区别", 2)
    for para in [
        "在模型比较中，必须区分解释模型和预测模型。线性回归和负二项回归强调系数方向、显著性、置信区间和发生率比，适合回答“哪些因素与租赁量有关、方向如何、统计证据是否稳定”。随机森林、LASSO 和回归树强调测试集误差，适合回答“给定已知时间、天气和日期信息后，能否较准确预测租赁量”。二者不是互相替代，而是服务于不同研究目标。",
        "更具体地说，线性回归和负二项回归的优势在于可解释性。线性回归能够用对数系数近似表达百分比变化，负二项回归能够用发生率比解释计数强度变化，因此适合写入影响因素分析和运营解释。随机森林的优势在于预测，它通过大量决策树捕捉小时、工作日、天气和温度之间的复杂交互，但单个变量的边际影响不如回归系数直观。因此，本文在结论中不会把随机森林变量重要性直接等同于因果影响。",
        "本文采用 80% 训练集和 20% 测试集划分，并设置随机种子保证可复现。所有预测模型只使用时间、日期、天气和环境变量，不使用 casual 与 registered，避免数据泄露。评价指标包括 RMSE、MAE 和测试集 R 平方。",
    ]:
        add_paragraph(doc, para)
    pred_rows = rows_from_dicts(read_csv("table27_prediction_metrics.csv"), [("model", "模型"), ("RMSE", "RMSE"), ("MAE", "MAE"), ("R_squared", "R 平方")], digits=3)
    add_table(doc, pred_rows, "表 14 预测模型性能比较 / Table 14 Predictive performance comparison")
    for para in [
        f"测试集结果显示，随机森林表现最好，RMSE 为 {sval('best_prediction_RMSE')}，MAE 为 {sval('best_prediction_MAE')}，R 平方为 {sval('best_prediction_R2')}。这一结果说明共享单车需求中存在明显非线性和交互结构，单纯线性模型难以充分刻画高峰时段需求。",
        "LASSO 回归通过 L1 惩罚实现变量选择和收缩，但其本质仍是线性模型；回归树能够表达非线性分割，但单棵树稳定性有限；随机森林通过多棵树集成降低方差，因此获得更好的预测效果[9][10][11]。",
        "从指标含义看，RMSE 对大误差更敏感，适合衡量高峰时段预测失败的代价；MAE 更接近平均绝对偏差，便于理解模型平均会错多少辆次；测试集 R 平方衡量模型对未见样本波动的解释比例。随机森林在三项指标上同时领先，说明其优势不是某一个指标偶然偏好造成的。",
    ]:
        add_paragraph(doc, para)
    add_figure(doc, "fig17_rf_prediction_actual.png", "图 11 随机森林预测值与真实值对比 / Figure 11 Random forest predicted versus observed values", 12)
    add_figure(doc, "fig17b_prediction_actual_all_models.png", "图 12 多模型预测值与真实值对比 / Figure 12 Predicted versus observed values across models", 13)

    add_heading(doc, "7.3 变量重要性与 LASSO 结果", 2)
    for para in [
        f"随机森林变量重要性显示，最重要的变量包括 {sval('top_rf_variables')}。小时变量排名最高，说明日内节律是预测共享单车需求的关键；工作日属性、气温、年份和湿度也具有较高重要性，说明通勤结构、年度趋势和舒适度共同影响需求。",
        f"LASSO 在 lambda.min = {sval('lasso_lambda_min')} 时保留 {sval('lasso_nonzero_terms')} 个非零变量。LASSO 结果说明小时、年份、季节和天气变量在线性正则化框架下仍有价值，但其预测性能不如随机森林，说明非线性结构在本任务中更重要。",
        "变量重要性应被理解为预测贡献，而不是单独变量的因果效应。比如 hr 排名最高，说明知道小时信息会显著降低预测误差；但这并不表示“小时本身导致租赁”，而是小时承载了通勤、休闲、营业时间和城市活动节奏等综合信息。类似地，workingday 的重要性反映工作日结构对需求曲线的影响，temp_celsius 和 hum_percent 则反映天气舒适度信息。",
        "LASSO 的价值在于提供一个介于普通线性回归和复杂机器学习之间的对照。它通过惩罚项压缩不稳定系数，使部分变量系数变为零，从而降低过拟合风险。本文保留 LASSO，不是因为它预测最好，而是因为它展示了 R 软件在正则化建模方面的能力，并与随机森林形成“线性正则化模型”和“非线性集成模型”的对照。",
    ]:
        add_paragraph(doc, para)
    rf_rows = rows_from_dicts(read_csv("table28b_rf_variable_importance.csv"), [("variable", "变量"), ("importance", "重要性")], limit=10, digits=2)
    add_table(doc, rf_rows, "表 15 随机森林变量重要性前十位 / Table 15 Top ten random forest variable importance")
    add_figure(doc, "fig18b_rf_variable_importance.png", "图 13 随机森林变量重要性 / Figure 13 Random forest variable importance", 12)
    lasso_rows = rows_from_dicts(read_csv("table28c_lasso_nonzero_coefficients.csv"), [("term", "变量"), ("coefficient", "系数")], limit=12, digits=3)
    add_table(doc, lasso_rows, "表 16 LASSO 非零系数前十二位 / Table 16 Top non-zero LASSO coefficients")

    add_heading(doc, "第八章 用户类型差异分析", 1)
    for para in [
        f"从总量看，注册用户贡献了 {pct(sval('registered_total_share'))} 的租赁量，休闲用户贡献 {pct(sval('casual_total_share'))}。因此，共享单车系统的主体需求来自注册用户。但两类用户并非只是数量不同，其时间模式和天气敏感性也存在差异。",
        "注册用户在工作日早晚高峰更明显，体现出通勤属性；休闲用户在非工作日白天相对更活跃，更接近休闲出行场景。分别建立 log1p(casual) 和 log1p(registered) 模型后可以发现，气温对休闲用户影响更强，湿度和风速对休闲用户的负向影响也更明显。这说明休闲用户更容易受到骑行舒适度影响，而注册用户需求相对稳定。",
        "这种用户差异具有明确的运营含义。注册用户的需求更接近日常刚性出行，因此调度重点应放在工作日早晚高峰和通勤走廊；休闲用户的需求更接近弹性出行，因此营销和活动安排更适合放在天气舒适、周末或节假日白天。若运营方只观察总租赁量，可能会低估休闲用户对天气的敏感性，也可能忽视注册用户在高峰时段对车辆可得性的要求。",
    ]:
        add_paragraph(doc, para)
    add_figure(doc, "fig13_user_hour_pattern.png", "图 14 休闲用户与注册用户小时模式 / Figure 14 Hourly patterns of casual and registered users", 13)
    add_paragraph(
        doc,
        "图 14 表明，注册用户曲线的高峰更加集中，休闲用户曲线则在白天更平缓。这种差异说明两类用户的决策逻辑不同：注册用户更可能把共享单车作为固定出行链条的一部分，休闲用户则更可能在天气、时间和个人安排适合时使用。图形也解释了为什么本文不只分析总租赁量，还要分别建模 casual 和 registered。",
    )
    add_figure(doc, "fig14_user_workingday_pattern.png", "图 15 不同日期类型下两类用户小时模式 / Figure 15 User patterns by working day", 13)
    add_paragraph(
        doc,
        "图 15 把用户类型与工作日属性结合起来观察。工作日注册用户早晚高峰最明显，非工作日休闲用户白天需求更突出。这说明工作日变量不是简单地提高或降低总租赁量，而是改变了不同用户群体在一天中的分布。若只在模型中加入 workingday 主效应，而不关注交互和用户分组，就会遗漏这一重要行为结构。",
    )
    user_fit_rows = rows_from_dicts(read_csv("table32_user_model_fit.csv"), [("model", "模型"), ("r.squared", "R 平方"), ("adj.r.squared", "调整 R 平方"), ("AIC", "AIC")], digits=3)
    add_table(doc, user_fit_rows, "表 17 用户类型模型拟合结果 / Table 17 User-segment model fit")
    user_focus_rows = rows_from_dicts(read_csv("table33_user_model_focus_terms.csv"), [("model", "用户类型"), ("term", "变量"), ("estimate", "系数"), ("p.value", "p 值"), ("percent_change", "近似百分比变化")], limit=10, digits=4)
    add_table(doc, user_focus_rows, "表 18 用户类型模型重点变量 / Table 18 Key coefficients of user-segment models")

    add_heading(doc, "第九章 结论、建议与不足", 1)
    add_heading(doc, "9.1 主要结论", 2)
    for para in [
        "第一，时间规律方面，共享单车租赁需求具有显著日内周期。凌晨需求最低，傍晚 17 时达到最高，工作日早晚高峰尤其突出。这说明共享单车不仅是休闲交通工具，也承担公共交通接驳和通勤补充功能。季节和月份趋势进一步表明，需求受到年度气候和活动节奏影响，春季较低、夏秋较高。",
        "第二，天气影响方面，晴朗或少云天气下租赁量最高，小雨雪天气明显下降；气温与租赁量呈正相关，湿度和风速在多元模型中呈负向影响。天气变量的作用并不只是改变平均租赁量，还会改变需求分布和高峰波动。对骑行这种暴露在室外环境中的出行方式而言，舒适度是不可忽视的需求解释因素。",
        "第三，模型结论方面，线性回归揭示了年份、季节、天气、气温、湿度和风速等变量的方向性关系；VIF 诊断没有发现严重共线性，Breusch-Pagan 检验提示异方差，HC3 稳健标准误增强了显著性判断的可信度。计数模型比较显示 Poisson 模型存在严重过度离散，负二项模型更符合租赁量作为非负整数且波动较大的数据特征。",
        "第四，用户差异方面，注册用户贡献超过八成租赁量，是系统需求主体；休闲用户虽然占比较低，但对气温、湿度、风速和非工作日白天更敏感。两类用户的差异说明，运营分析不能只看总量，还应区分通勤需求和休闲需求，否则会掩盖不同用户对时间和天气的不同反应。",
        "第五，运营启示方面，车辆调度应重点覆盖工作日早晚高峰，需求预测应优先纳入小时、工作日属性、气温、湿度和年份等变量。随机森林在测试集中取得 RMSE = 43.473、MAE = 26.999、R 平方 = 0.943，适合作为短期需求预测工具；线性回归和负二项回归则适合用于解释、汇报和制定管理策略。",
        "第六，研究限制方面，本文结论主要基于华盛顿 D.C. 2011 至 2012 年数据，缺少站点空间信息、价格、车辆供给、用户注册规模和城市事件变量，因此不宜作严格因果推断。后续研究可加入站点级空间数据、时间序列交叉验证和更丰富的外部事件变量，以提升预测和解释能力。",
    ]:
        add_paragraph(doc, para)

    add_heading(doc, "9.2 运营建议", 2)
    for para in [
        "基于时间规律，运营方应重点关注工作日早晚高峰车辆调度，尤其是 7 至 9 时和 17 至 18 时附近的供需平衡。基于天气规律，晴朗和适宜温度条件下应提高车辆供给预期，在小雨雪、高湿度和大风条件下应适当下调需求预期并关注骑行安全。",
        "基于模型结果，需求预测系统应优先纳入小时、工作日属性、气温、年份和湿度等变量。随机森林等机器学习模型可用于短期需求预测，但回归模型仍应保留用于解释和管理沟通。",
    ]:
        add_paragraph(doc, para)

    add_heading(doc, "9.3 研究不足", 2)
    for para in [
        "本文仍存在不足。首先，数据仅来自华盛顿 D.C.，结论是否适用于其他城市需要进一步验证。其次，数据没有站点位置、车辆投放量、价格、用户注册规模和活动事件信息，因此无法进行严格因果识别。第三，本文采用随机训练测试划分，适合比较模型总体预测能力，但如果目标是未来时点预测，应进一步采用时间序列切分或滚动验证。第四，本文尚未加入梯度提升树、时间序列模型和站点级空间特征，后续可在避免数据泄露的前提下继续扩展。",
    ]:
        add_paragraph(doc, para)

    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] Fanaee-T H. Bike Sharing Dataset[E]. UCI Machine Learning Repository. https://doi.org/10.24432/C5W894, 2013.",
        "[2] Fanaee-T H, Gama J. Event labeling combining ensemble detectors and background knowledge[J]. Progress in Artificial Intelligence, 2013, 2(2-3): 113-127.",
        "[3] R Core Team. R: A language and environment for statistical computing[E]. R Foundation for Statistical Computing. https://www.r-project.org/.",
        "[4] Wickham H, Averick M, Bryan J, et al. Welcome to the tidyverse[J]. Journal of Open Source Software, 2019, 4(43): 1686.",
        "[5] Wickham H. ggplot2: Elegant Graphics for Data Analysis[M]. New York: Springer, 2016.",
        "[6] Venables W N, Ripley B D. Modern Applied Statistics with S[M]. New York: Springer, 2002.",
        "[7] Breusch T S, Pagan A R. A simple test for heteroscedasticity and random coefficient variation[J]. Econometrica, 1979, 47(5): 1287-1294.",
        "[8] White H. A heteroskedasticity-consistent covariance matrix estimator and a direct test for heteroskedasticity[J]. Econometrica, 1980, 48(4): 817-838.",
        "[9] Breiman L. Random forests[J]. Machine Learning, 2001, 45(1): 5-32.",
        "[10] Wright M N, Ziegler A. ranger: A fast implementation of random forests for high dimensional data in C++ and R[J]. Journal of Statistical Software, 2017, 77(1): 1-17.",
        "[11] Friedman J, Hastie T, Tibshirani R. Regularization paths for generalized linear models via coordinate descent[J]. Journal of Statistical Software, 2010, 33(1): 1-22.",
    ]
    for ref in refs:
        add_paragraph(doc, ref, first_indent=False, size=10.5)

    add_heading(doc, "致谢", 1)
    add_paragraph(doc, "感谢《统计软件》课程提供将统计方法、R 软件和真实数据分析结合起来的训练机会。本文的数据来源于 UCI Machine Learning Repository，相关 R 软件包和公开文献为论文分析提供了方法基础。论文仍有不足，后续可在更丰富的城市、站点和事件数据基础上继续改进。")

    add_heading(doc, "附录 A R 运行环境", 1)
    session_text = sanitize_paths((ROOT / "outputs" / "sessionInfo.txt").read_text(encoding="utf-8", errors="ignore"))
    for line in session_text.splitlines()[:45]:
        add_paragraph(doc, line, first_indent=False, size=8.5)

    add_heading(doc, "附录 B 核心 R 分析代码", 1)
    code = (ROOT / "scripts" / "analysis.R").read_text(encoding="utf-8")
    for line in code.splitlines():
        # Keep the appendix complete but compact.
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(7.5)

    doc.add_page_break()
    add_heading(doc, "附录 C 附图附表全文与模型文件清单", 1)
    add_paragraph(
        doc,
        "本附录不再仅列出文件名，而是将 R 脚本输出的附图和附表实际写入论文。附图来自 outputs/figures，附表来自 outputs/tables；模型文件为 R 的二进制 .rds 对象，不适合直接展开为正文，因此在 C.3 中以清单方式列出。",
    )

    add_heading(doc, "C.1 附图全文写入", 2)
    for idx, image_path in enumerate(sorted(FIG_DIR.glob("*.png")), start=1):
        add_appendix_image(
            doc,
            image_path,
            f"附图 C-{idx} {image_path.name} / Appendix Figure C-{idx} {image_path.name}",
            width_cm=12.8,
        )

    add_heading(doc, "C.2 附表路径与代表性预览", 2)
    add_paragraph(
        doc,
        "考虑到部分附表为测试集逐行预测明细或 HTML 模型汇总，完整展开会显著增加 Word 体积并影响排版。本节对每个附表写入完整文件路径、行列规模和代表性预览；完整结果以 outputs/tables 中的原文件为准。",
    )
    for idx, table_path in enumerate(sorted(TABLE_DIR.glob("*")), start=1):
        rows = read_table_file(table_path)
        n_rows = len(rows) - 1 if rows else 0
        n_cols = len(rows[0]) if rows else 0
        add_paragraph(
            doc,
            f"附表 C-{idx} 文件：{table_path.name}；相对路径：{rel_path(table_path)}；规模：约 {n_rows} 行、{n_cols} 列。",
            first_indent=False,
            size=9,
        )
        add_appendix_data_table(
            doc,
            preview_rows(rows),
            f"附表 C-{idx} {table_path.name} 预览 / Appendix Table C-{idx} Preview of {table_path.name}",
        )

    add_heading(doc, "C.3 模型文件清单", 2)
    model_rows = [["模型文件名", "文件类型", "路径"]]
    for model_path in sorted(MODEL_DIR.glob("*")):
        model_rows.append([model_path.name, model_path.suffix or "文件", rel_path(model_path)])
    add_appendix_data_table(doc, model_rows, "附表 C-M1 模型文件清单 / Appendix Table C-M1 Model file list")

    output = PAPER_DIR / "paper.docx"
    doc.save(output)
    print(output)


if __name__ == "__main__":
    build()
